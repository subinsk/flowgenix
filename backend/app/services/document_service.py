import os
import aiofiles
from typing import List, Optional, Dict, Any
from sqlalchemy.orm import Session
from fastapi import UploadFile
import fitz  # PyMuPDF
import chromadb
from docx import Document as DocxDocument

from app.models.document import Document
from app.schemas.document import DocumentCreate
from app.services.ai_service import AIService
from app.services.api_key_service import ApiKeyService
from app.core.config import settings


class DocumentService:
    def __init__(self, db: Session = None):
        self.db = db
        self.upload_dir = settings.upload_dir
        
        # Initialize ChromaDB in embedded mode
        self.chroma_client = chromadb.Client(chromadb.config.Settings(
            persist_directory=settings.chroma_persist_directory,
            anonymized_telemetry=False
        ))
        
        self.collection_name = "documents"
        self.api_key_service = ApiKeyService(db) if db else None
        
        # Ensure upload directory exists
        os.makedirs(self.upload_dir, exist_ok=True)
        
        # Ensure ChromaDB persist directory exists
        os.makedirs(settings.chroma_persist_directory, exist_ok=True)

    async def upload_document(self, file: UploadFile, user_id: str, embedding_model: str = "text-embedding-ada-002", api_key: str = None) -> Document:
        """Upload and process document"""
        # Validate file type
        if not self._is_valid_file_type(file.filename):
            raise ValueError("Invalid file type")

        # Save file
        file_path = os.path.join(self.upload_dir, f"{user_id}_{file.filename}")
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        # Create database record
        document_data = DocumentCreate(
            filename=file.filename,
            content_type=file.content_type,
            file_size=len(content),
            file_path=file_path
        )

        db_document = Document(
            filename=document_data.filename,
            content_type=document_data.content_type,
            file_size=document_data.file_size,
            file_path=document_data.file_path,
            user_id=user_id
        )

        if self.db:
            self.db.add(db_document)
            self.db.commit()
            self.db.refresh(db_document)

        # Process document asynchronously with specified model and API key
        await self._process_document(db_document, embedding_model, api_key)

        return db_document

    async def extract_text(self, document_id: str) -> str:
        """Extract text from document"""
        if not self.db:
            return ""

        document = self.db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise ValueError("Document not found")

        return await self._extract_text_from_file(document.file_path)

    async def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text from file based on type"""
        if file_path.endswith('.pdf'):
            return self._extract_pdf_text(file_path)
        elif file_path.endswith('.txt'):
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        elif file_path.endswith('.docx'):
            return self._extract_docx_text(file_path)
        else:
            return "Unsupported file type for text extraction"

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            return f"Error extracting PDF text: {str(e)}"

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            return f"Error extracting DOCX text: {str(e)}"

    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content using PyMuPDF"""
        try:
            import tempfile
            import io
            
            # Check if content is valid
            if not content or len(content) == 0:
                return "Error: PDF content is empty"
                
            # Try to open directly from bytes first
            try:
                doc = fitz.open(stream=content, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                
                if not text.strip():
                    return "Error: PDF appears to be empty or contains no extractable text"
                    
                return text
            except Exception as direct_error:
                print(f"Direct PDF processing failed: {direct_error}")
                
                # Fallback: Use temporary file
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                    temp_file.write(content)
                    temp_file.flush()
                    
                    doc = fitz.open(temp_file.name)
                    text = ""
                    for page in doc:
                        text += page.get_text()
                    doc.close()
                    
                    # Clean up temp file
                    os.unlink(temp_file.name)
                    
                    if not text.strip():
                        return "Error: PDF appears to be empty or contains no extractable text"
                        
                    return text
                    
        except Exception as e:
            print(f"PDF extraction error: {str(e)}")
            return f"Error extracting PDF text: {str(e)}"

    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX content using python-docx"""
        try:
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                doc = DocxDocument(temp_file.name)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                # Clean up temp file
                os.unlink(temp_file.name)
                return text
        except Exception as e:
            return f"Error extracting DOCX text: {str(e)}"

    async def _process_document(self, document: Document, embedding_model: str = "text-embedding-ada-002", api_key: str = None):
        """Process document: extract text and generate embeddings with selected model and key"""
        try:
            # Extract text
            text = await self._extract_text_from_file(document.file_path)
            
            # Get API key if not provided
            final_api_key = api_key
            if not final_api_key and self.api_key_service and document.user_id:
                # Determine which API key to use based on embedding model
                if embedding_model == "all-MiniLM-L6-v2":
                    final_api_key = self.api_key_service.get_decrypted_api_key(str(document.user_id), "huggingface")
                else:
                    final_api_key = self.api_key_service.get_decrypted_api_key(str(document.user_id), "openai")
            
            # Generate embeddings with model and key
            ai_service = AIService()
            embeddings = await ai_service.generate_embeddings(text, model=embedding_model, api_key=final_api_key)
            if embeddings:
                # Store in ChromaDB
                collection = self._get_or_create_collection()
                collection.add(
                    documents=[text],
                    metadatas=[{
                        "document_id": str(document.id),
                        "filename": document.filename,
                        "user_id": str(document.user_id),
                        "workflow_id": str(document.workflow_id) if document.workflow_id else None
                    }],
                    ids=[str(document.id)],
                    embeddings=[embeddings]
                )
                # Mark as processed only if embeddings were successfully generated and stored
                if self.db:
                    document.processed = True
                    self.db.commit()
                print(f"Successfully processed document {document.id} with {embedding_model}")
            else:
                # Don't mark as processed if embedding generation failed
                print(f"Failed to generate embeddings for document {document.id} using {embedding_model}")
                if embedding_model == "all-MiniLM-L6-v2" and final_api_key:
                    print("Note: HuggingFace sentence-transformers models currently have issues with the Inference API")
                    print("Consider using OpenAI embeddings (text-embedding-ada-002) as an alternative")
                elif not final_api_key:
                    print(f"No API key available for {embedding_model}")
        except Exception as e:
            print(f"Error processing document {document.id}: {str(e)}")
            # Don't mark as processed if there was an error

    async def search_documents_by_user(self, query: str, user_id: str, limit: int = 5, embedding_model: str = "all-MiniLM-L6-v2") -> List[Dict[str, Any]]:
        """Search documents using vector similarity"""
        try:
            # Generate query embedding using the same model as documents
            ai_service = AIService()
            
            # Get API key if needed
            api_key = None
            if self.api_key_service and embedding_model == "all-MiniLM-L6-v2":
                api_key = self.api_key_service.get_decrypted_api_key(user_id, "huggingface")
            elif self.api_key_service and embedding_model.startswith("text-embedding"):
                api_key = self.api_key_service.get_decrypted_api_key(user_id, "openai")
            
            query_embedding = await ai_service.generate_embeddings(query, model=embedding_model, api_key=api_key)
            
            if not query_embedding:
                return []

            # Search in ChromaDB
            collection = self._get_or_create_collection()
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=limit,
                where={"user_id": user_id}
            )

            # Format results
            search_results = []
            for i, doc_id in enumerate(results['ids'][0]):
                search_results.append({
                    "document_id": str(doc_id),
                    "content": results['documents'][0][i][:500] + "...",
                    "similarity": 1 - results['distances'][0][i],
                    "metadata": results['metadatas'][0][i]
                })

            return search_results

        except Exception as e:
            print(f"Error searching documents: {str(e)}")
            return []

    async def search_documents(self, query: str, document_ids: List[str] = None, limit: int = 5) -> List[Dict[str, Any]]:
        """Search documents for workflow execution"""
        try:
            # Generate embedding for query
            ai_service = AIService()
            query_embedding = await ai_service.generate_embeddings(query)
            
            if not query_embedding:
                return []

            # Search in ChromaDB
            collection = self._get_or_create_collection()
            
            # Build where clause
            where_clause = {}
            if document_ids:
                where_clause["document_id"] = {"$in": document_ids}
            
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=limit,
                where=where_clause if where_clause else None
            )

            # Format results for workflow
            search_results = []
            for i, doc_id in enumerate(results['ids'][0]):
                search_results.append({
                    "id": doc_id,
                    "content": results['documents'][0][i],
                    "similarity": 1 - results['distances'][0][i],
                    "metadata": results['metadatas'][0][i] or {}
                })

            return search_results

        except Exception as e:
            print(f"Error searching documents for workflow: {str(e)}")
            return []

    async def process_document(self, file: UploadFile, workflow_id: str, user_id: str) -> Dict[str, Any]:
        """Process and store document for workflow"""
        try:
            # Save and process document
            content = await file.read()
            
            print(f"Processing document: {file.filename}, size: {len(content)} bytes, type: {file.content_type}")
            
            # Check if content is valid
            if not content or len(content) == 0:
                raise ValueError("Document content is empty")
            
            # Extract text content
            text_content = ""
            if file.content_type == "application/pdf":
                text_content = await self._extract_pdf_text(content)
            elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text_content = await self._extract_docx_text(content)
            elif file.content_type.startswith("text/"):
                text_content = content.decode('utf-8')
            else:
                raise ValueError(f"Unsupported file type: {file.content_type}")
            
            print(f"Extracted text length: {len(text_content)}")
            
            if not text_content.strip():
                raise ValueError("No text could be extracted from the document")
            
            # Save to database first
            from app.models.document import Document
            import os
            from app.core.config import settings
            import aiofiles
            
            # Save file to disk
            file_path = os.path.join(settings.upload_dir, f"{user_id}_{file.filename}")
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(content)
            
            # Create DB record
            db_document = Document(
                filename=file.filename,
                content_type=file.content_type,
                file_size=len(content),
                file_path=file_path,
                user_id=user_id,
                workflow_id=workflow_id,
                processed=False
            )
            
            if self.db:
                self.db.add(db_document)
                self.db.commit()
                self.db.refresh(db_document)
            
            # Generate embeddings with user's API keys
            print(f"DEBUG: Generating embeddings for document: {file.filename}")
            print(f"DEBUG: Text content length: {len(text_content)}")
            ai_service = AIService()
            
            # Try to get user's API keys for embeddings
            embeddings = None
            api_key = None
            embedding_model = "all-MiniLM-L6-v2"  # Default to HuggingFace free model
            
            if self.api_key_service and user_id:
                # Try HuggingFace first (free tier)
                api_key = self.api_key_service.get_decrypted_api_key(str(user_id), "huggingface")
                if api_key:
                    print(f"DEBUG: Using HuggingFace API key for embeddings")
                    embeddings = await ai_service.generate_embeddings(text_content, model=embedding_model, api_key=api_key)
                else:
                    # Try OpenAI as fallback
                    api_key = self.api_key_service.get_decrypted_api_key(str(user_id), "openai")
                    if api_key:
                        print(f"DEBUG: Using OpenAI API key for embeddings")
                        embedding_model = "text-embedding-ada-002"
                        embeddings = await ai_service.generate_embeddings(text_content, model=embedding_model, api_key=api_key)
                    else:
                        print("DEBUG: No API keys found for user - embeddings will not be generated")
            else:
                print("DEBUG: API key service not available - embeddings will not be generated")
            
            print(f"DEBUG: Embeddings generated: {bool(embeddings)}")
            if embeddings:
                print(f"DEBUG: Embedding dimensions: {len(embeddings)}")
            else:
                print("WARNING: No embeddings generated - document will not be searchable")
            
            doc_id = f"{workflow_id}_{file.filename}_{hash(text_content)}"
            
            if embeddings:
                print(f"Generated embeddings, length: {len(embeddings)}")
                
                # Store in ChromaDB with error handling for dimension mismatch
                try:
                    collection = self._get_or_create_collection()
                    
                    collection.add(
                        documents=[text_content],
                        embeddings=[embeddings],
                        metadatas=[{
                            "filename": file.filename,
                            "workflow_id": workflow_id,
                            "user_id": user_id,
                            "content_type": file.content_type,
                            "size": len(content),
                            "doc_id": str(db_document.id) if hasattr(db_document, 'id') else doc_id
                        }],
                        ids=[doc_id]
                    )
                    print(f"✓ Document stored in ChromaDB with embeddings: {doc_id}")
                except Exception as e:
                    if "dimension" in str(e).lower():
                        print(f"Dimension mismatch detected: {str(e)}")
                        print("Recreating collection with correct dimensions...")
                        collection = self._get_or_create_collection(force_recreate=True)
                        collection.add(
                            documents=[text_content],
                            embeddings=[embeddings],
                            metadatas=[{
                                "filename": file.filename,
                                "workflow_id": workflow_id,
                                "user_id": user_id,
                                "content_type": file.content_type,
                                "size": len(content),
                                "doc_id": str(db_document.id) if hasattr(db_document, 'id') else doc_id
                            }],
                            ids=[doc_id]
                        )
                        print(f"✓ Document stored in ChromaDB with embeddings after recreation: {doc_id}")
                    else:
                        raise e
            else:
                print("WARNING: No embeddings available - storing document without embeddings")
                # Store document without embeddings (text-only for fallback retrieval)
                try:
                    collection = self._get_or_create_collection()
                    collection.add(
                        documents=[text_content],
                        metadatas=[{
                            "filename": file.filename,
                            "workflow_id": workflow_id,
                            "user_id": user_id,
                            "content_type": file.content_type,
                            "size": len(content),
                            "doc_id": str(db_document.id) if hasattr(db_document, 'id') else doc_id,
                            "no_embeddings": True  # Flag for fallback retrieval
                        }],
                        ids=[doc_id]
                    )
                    print(f"✓ Document stored in ChromaDB without embeddings: {doc_id}")
                except Exception as e:
                    if "dimension" in str(e).lower():
                        print(f"Dimension mismatch detected for no-embeddings storage: {str(e)}")
                        print("Recreating collection...")
                        collection = self._get_or_create_collection(force_recreate=True)
                        collection.add(
                            documents=[text_content],
                            metadatas=[{
                                "filename": file.filename,
                                "workflow_id": workflow_id,
                                "user_id": user_id,
                                "content_type": file.content_type,
                                "size": len(content),
                                "doc_id": str(db_document.id) if hasattr(db_document, 'id') else doc_id,
                                "no_embeddings": True  # Flag for fallback retrieval
                            }],
                            ids=[doc_id]
                        )
                        print(f"✓ Document stored in ChromaDB without embeddings after recreation: {doc_id}")
                    else:
                        raise e
                
            # Mark as processed
            if self.db and hasattr(db_document, 'id'):
                db_document.processed = True
                self.db.commit()
                
            print(f"Document stored in ChromaDB with ID: {doc_id}")
            
            return {
                "doc_id": doc_id,
                "filename": file.filename,
                "size": len(content),
                "content_type": file.content_type,
                "processed": bool(embeddings),
                "embeddings_count": len(embeddings) if embeddings else 0,
                "text_length": len(text_content),
                "message": f"Successfully processed {file.filename}"
            }
            
        except Exception as e:
            print(f"Document processing error: {str(e)}")
            raise ValueError(f"Failed to process document: {str(e)}")

    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            raise ValueError(f"Failed to extract PDF text: {str(e)}")

    def _get_or_create_collection(self, force_recreate=False):
        """Get or create ChromaDB collection with proper embedding configuration"""
        try:
            if force_recreate:
                # Delete existing collection if it exists
                try:
                    self.chroma_client.delete_collection(self.collection_name)
                    print(f"Deleted existing collection: {self.collection_name}")
                except:
                    pass
            
            # Try to get existing collection first
            collection = self.chroma_client.get_collection(self.collection_name)
            print(f"Using existing collection: {self.collection_name}")
            return collection
        except:
            # Create new collection - let ChromaDB use its default embedding function
            # This will use all-MiniLM-L6-v2 which produces 384-dimensional embeddings
            print(f"Creating new collection: {self.collection_name}")
            return self.chroma_client.create_collection(self.collection_name)

    def _is_valid_file_type(self, filename: str) -> bool:
        """Check if file type is allowed"""
        if not filename:
            return False
        
        file_ext = os.path.splitext(filename)[1].lower()
        return file_ext in settings.allowed_file_types

    def get_user_documents(self, user_id: str, skip: int = 0, limit: int = 100) -> List[Document]:
        """Get all documents for a user"""
        if not self.db:
            return []
        
        return self.db.query(Document).filter(
            Document.user_id == user_id
        ).offset(skip).limit(limit).all()
